#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate complex test document for testing HybridChunker and document processing pipeline.

This script creates a realistic vehicle service report with:
- Multiple heading levels with bold formatting
- Line breaks and paragraphs
- Long tables (testing chunking boundaries)
- Nested lists
- Mixed formatting (bold, italic)
- Image placeholder (for OCR testing)

Usage:
    python generate_complex_test_document.py

Output:
    tests/test_data/Vehicle_Service_Report_Toyota_Camry_2023.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os

def add_image_placeholder(doc, text="VIN: WF0XXXXXXXXXXXXX1"):
    """
    Add a text box that simulates an image with text (for OCR testing).
    In real use, you would replace this with an actual image.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a styled text box to simulate image
    run = paragraph.add_run(f"\n[IMAGE: Vehicle VIN Plate]\n{text}\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True

    # Add note
    note = doc.add_paragraph()
    note_run = note.add_run("(Note: Replace this with actual VIN plate image for OCR testing)")
    note_run.font.size = Pt(8)
    note_run.italic = True
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_complex_test_document():
    """Create a complex DOCX document for testing."""

    # Create output directory
    output_dir = Path(__file__).parent / "test_data"
    output_dir.mkdir(exist_ok=True)

    output_path = output_dir / "Vehicle_Service_Report_Toyota_Camry_2023.docx"

    # Create document
    doc = Document()

    # =====================================================================
    # PAGE 1: Title and Vehicle Information
    # =====================================================================

    # Main title (Heading 1, bold, large)
    title = doc.add_heading('VEHICLE SERVICE REPORT', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Comprehensive Maintenance Report')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Line break

    # Section: Vehicle Information
    doc.add_heading('Vehicle Information', level=2)

    # Mixed formatting paragraph
    p1 = doc.add_paragraph()
    p1.add_run('Registration Number: ').bold = False
    p1.add_run('191-D-12345').bold = True

    p2 = doc.add_paragraph()
    p2.add_run('VIN: ').bold = False
    p2.add_run('WF0XXXXXXXXXXXXX1').bold = True

    p3 = doc.add_paragraph()
    p3.add_run('Make/Model: ').bold = False
    p3.add_run('Toyota Camry 2023').bold = True

    p4 = doc.add_paragraph()
    p4.add_run('Service Date: ')
    p4.add_run('15 January 2025').italic = True

    p5 = doc.add_paragraph()
    p5.add_run('Current Mileage: ').bold = False
    p5.add_run('45,230 km').bold = True

    doc.add_paragraph()  # Line break

    # Owner information
    doc.add_heading('Owner Information', level=2)

    owner_p = doc.add_paragraph()
    owner_p.add_run('Owner: John Murphy\n')
    owner_p.add_run('Contact: +353 86 123 4567\n')
    owner_p.add_run('Email: john.murphy@example.ie\n')
    owner_p.add_run('Address: 15 Main Street, Dublin 2, Ireland')

    doc.add_paragraph()  # Line break

    # Service Overview
    doc.add_heading('Service Overview', level=2)

    overview = doc.add_paragraph(
        'This comprehensive service report documents all maintenance activities performed '
        'on the vehicle during the scheduled 45,000 km service interval. '
        'All work was completed in accordance with Toyota manufacturer specifications '
        'and European Union vehicle safety standards.'
    )

    doc.add_paragraph()  # Line break

    # Bulleted list: Services Performed
    doc.add_heading('Services Performed', level=3)

    services = [
        'Full engine oil and filter replacement',
        'Brake system inspection and pad replacement',
        'Tire rotation and pressure adjustment',
        'Air conditioning system service',
        'Battery health check and terminal cleaning',
        'Fluid level inspection (coolant, brake, power steering)',
        'Lighting system check (all bulbs and indicators)',
        'Windscreen wiper blade replacement'
    ]

    for service in services:
        doc.add_paragraph(service, style='List Bullet')

    doc.add_page_break()

    # =====================================================================
    # PAGE 2: Long Service History Table (CRITICAL for chunking test)
    # =====================================================================

    doc.add_heading('Complete Service History', level=2)

    intro = doc.add_paragraph(
        'The following table shows the complete maintenance history for this vehicle. '
        'This data is critical for warranty validation and resale value assessment.'
    )

    # Create large table (15 rows + header)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Date', 'Mileage', 'Service Type', 'Parts Replaced', 'Cost (€)', 'Technician']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Bold header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Service history data (15 entries)
    service_data = [
        ['15/01/2025', '45,230 km', 'Full Service', 'Oil filter, Air filter, Wiper blades', '€285.00', 'John Smith'],
        ['10/12/2024', '44,100 km', 'Tire Service', 'Tire rotation, Pressure adjustment', '€45.00', 'Mary Jones'],
        ['05/11/2024', '43,000 km', 'Brake Inspection', 'Brake pads (front), Brake fluid', '€320.00', 'John Smith'],
        ['20/10/2024', '42,500 km', 'NCT Preparation', 'NCT inspection, Emissions test', '€95.00', 'Mary Jones'],
        ['15/09/2024', '41,800 km', 'Oil Change', 'Engine oil, Oil filter', '€120.00', 'Tom Brown'],
        ['10/08/2024', '40,200 km', 'Air Conditioning', 'AC service, Refrigerant refill', '€150.00', 'John Smith'],
        ['05/07/2024', '39,100 km', 'Tire Replacement', 'Front tires (2x Michelin)', '€380.00', 'Mary Jones'],
        ['15/06/2024', '38,500 km', 'Battery Service', 'Battery replacement, Terminal cleaning', '€180.00', 'Tom Brown'],
        ['10/05/2024', '37,200 km', 'Brake Service', 'Brake pads (rear), Brake disc machining', '€340.00', 'John Smith'],
        ['05/04/2024', '36,000 km', 'Major Service', 'Oil, filters, spark plugs, coolant', '€450.00', 'John Smith'],
        ['15/03/2024', '35,100 km', 'Suspension Check', 'Suspension inspection, Alignment', '€110.00', 'Mary Jones'],
        ['10/02/2024', '34,500 km', 'Light Repair', 'Headlight bulb replacement', '€45.00', 'Tom Brown'],
        ['05/01/2024', '33,800 km', 'Oil Change', 'Engine oil, Oil filter', '€120.00', 'John Smith'],
        ['20/12/2023', '32,500 km', 'Winter Check', 'Antifreeze top-up, Battery check', '€65.00', 'Mary Jones'],
        ['15/11/2023', '31,200 km', 'Tire Service', 'Winter tire installation', '€80.00', 'Tom Brown'],
    ]

    # Add data rows
    for row_data in service_data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = value

    # Add total row
    total_row = table.add_row().cells
    total_row[0].text = 'TOTAL'
    for paragraph in total_row[0].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
    total_row[4].text = '€2,785.00'
    for paragraph in total_row[4].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

    doc.add_paragraph()  # Line break

    # Note about table
    note = doc.add_paragraph()
    note_run = note.add_run('Note: ')
    note_run.bold = True
    note.add_run('This table should remain intact during chunking. '
                'HybridChunker should not split table rows across chunks.')

    doc.add_page_break()

    # =====================================================================
    # PAGE 3: Image with Text (OCR Testing)
    # =====================================================================

    doc.add_heading('Vehicle Identification', level=2)

    p = doc.add_paragraph(
        'The Vehicle Identification Number (VIN) is located on the dashboard '
        'visible through the windscreen, and also on the vehicle registration certificate.'
    )

    doc.add_paragraph()  # Line break

    # Add image placeholder
    add_image_placeholder(doc, "VIN: WF0XXXXXXXXXXXXX1")

    doc.add_paragraph()  # Line break

    # VIN breakdown
    doc.add_heading('VIN Breakdown', level=3)

    vin_list = [
        'WF0 - World Manufacturer Identifier (Toyota Europe)',
        'XXXXXXXX - Vehicle Descriptor Section',
        'XXX1 - Vehicle Identifier Section (Sequential Number)'
    ]

    for item in vin_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =====================================================================
    # PAGE 4: Nested Structure and Cost Breakdown
    # =====================================================================

    doc.add_heading('Detailed Parts and Labor Breakdown', level=2)

    # Nested sections (testing heading hierarchy)
    doc.add_heading('Engine Components', level=3)

    doc.add_heading('Oil System', level=4)
    p1 = doc.add_paragraph('Engine oil (5W-30 synthetic, 4.5L): ')
    p1.add_run('€45.00').bold = True

    p2 = doc.add_paragraph('Oil filter (OEM Toyota): ')
    p2.add_run('€18.00').bold = True

    doc.add_heading('Air Intake', level=4)
    p3 = doc.add_paragraph('Air filter (high-performance): ')
    p3.add_run('€25.00').bold = True

    p4 = doc.add_paragraph('Cabin air filter (activated carbon): ')
    p4.add_run('€22.00').bold = True

    # Brake System
    doc.add_heading('Brake System', level=3)

    doc.add_heading('Front Brakes', level=4)
    p5 = doc.add_paragraph('Brake pads (ceramic, set of 4): ')
    p5.add_run('€120.00').bold = True

    p6 = doc.add_paragraph('Brake disc machining (front): ')
    p6.add_run('€60.00').bold = True

    doc.add_heading('Brake Fluids', level=4)
    p7 = doc.add_paragraph('DOT 4 brake fluid (1L): ')
    p7.add_run('€15.00').bold = True

    doc.add_paragraph()  # Line break

    # Cost breakdown table
    doc.add_heading('Service Cost Summary', level=2)

    cost_table = doc.add_table(rows=1, cols=2)
    cost_table.style = 'Light List Accent 1'

    # Header
    cost_header = cost_table.rows[0].cells
    cost_header[0].text = 'Item'
    cost_header[1].text = 'Amount (€)'
    for cell in cost_header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Cost items
    cost_items = [
        ['Parts Total', '€305.00'],
        ['Labor (3.5 hours @ €65/hr)', '€227.50'],
        ['Subtotal', '€532.50'],
        ['VAT (23%)', '€122.48'],
        ['TOTAL', '€654.98']
    ]

    for item, amount in cost_items:
        row = cost_table.add_row().cells
        row[0].text = item
        row[1].text = amount

        # Bold the TOTAL row
        if item == 'TOTAL':
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)

    doc.add_paragraph()  # Line break

    # Final summary
    final_para = doc.add_paragraph()
    final_para.add_run('Total Service Cost: ')
    total_run = final_para.add_run('€654.98')
    total_run.bold = True
    total_run.font.size = Pt(14)
    total_run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph()  # Line break

    # Payment and warranty info
    doc.add_heading('Payment and Warranty', level=2)

    payment_para = doc.add_paragraph(
        'Payment received in full on 15/01/2025 via credit card. '
    )
    warranty_run = payment_para.add_run('All parts and labor are covered by a 12-month warranty ')
    warranty_run.bold = True
    payment_para.add_run(
        'in accordance with EU consumer protection regulations. '
        'This warranty covers defects in materials and workmanship.'
    )

    doc.add_paragraph()  # Line break

    # Signature section
    doc.add_heading('Service Certification', level=3)

    sig_para = doc.add_paragraph()
    sig_para.add_run('Certified by: ')
    sig_para.add_run('John Smith, Senior Technician').bold = True
    sig_para.add_run('\nCertification Number: ')
    sig_para.add_run('TC-2025-001234').bold = True
    sig_para.add_run('\nDate: ')
    sig_para.add_run('15 January 2025').italic = True

    # Save document
    doc.save(str(output_path))

    print("=" * 70)
    print("[OK] COMPLEX TEST DOCUMENT CREATED SUCCESSFULLY")
    print("=" * 70)
    print(f"Location: {output_path}")
    print(f"File size: {output_path.stat().st_size / 1024:.1f} KB")
    print("\nDocument structure:")
    print("  - 4 pages")
    print("  - 3 heading levels (H1, H2, H3, H4)")
    print("  - 1 large table (16 rows, 6 columns)")
    print("  - 1 small table (cost breakdown)")
    print("  - Mixed formatting (bold, italic, colors)")
    print("  - Bulleted and numbered lists")
    print("  - Image placeholder (for OCR testing)")
    print("  - Nested sections (heading hierarchy)")
    print("\nNext steps:")
    print("  1. Open the DOCX file in Word")
    print("  2. (Optional) Replace image placeholder with real VIN plate image")
    print("  3. Save as PDF: 'Vehicle_Service_Report_Toyota_Camry_2023.pdf'")
    print("  4. Run test: python tests/test_complex_document_processing.py")
    print("=" * 70)

    return output_path

if __name__ == "__main__":
    create_complex_test_document()
